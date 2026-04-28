# -*- coding: utf-8 -*-
"""生成面向终端用户的 Word 版使用说明书（通俗表述，可重复运行）。"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, Cm
from docx.enum.style import WD_STYLE_TYPE


def set_doc_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    doc.add_paragraph()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
    else:
        run.font.size = Pt(14)
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    doc.add_paragraph()


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(11)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def add_table_two_col(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows) + 1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "内容"
    hdr[1].text = "说明"
    for i, (a, b) in enumerate(rows, start=1):
        table.rows[i].cells[0].text = a
        table.rows[i].cells[1].text = b
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "宋体"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                    run.font.size = Pt(10.5)
    doc.add_paragraph()


def build() -> Path:
    out_dir = Path(__file__).resolve().parent.parent
    # 使用 Unicode 转义，避免源文件编码导致 Windows 下文件名乱码
    out_path = out_dir / "\u706b\u6e90\u7075\u611f\u706b\u82b1-\u7528\u6237\u4f7f\u7528\u8bf4\u660e\u4e66.docx"

    doc = Document()
    set_doc_defaults(doc)

    add_title(doc, "火源灵感火花 · 小程序使用说明书")

    add_body(
        doc,
        "欢迎使用「火源灵感火花」小程序。本程序帮您用人工智能辅助做文案、聊创作、管理自己的账号形象（IP）。"
        "屏幕最下方有三个主按钮：发现、IP 操作台、我的，点一下即可切换。",
    )
    add_body(
        doc,
        "说明：小程序里有时显示「火源文案」，与「火源灵感火花」指同一款产品。文中所说「AI 助手」即为您服务的智能程序，会按您选的入口帮您写稿、对话。",
    )

    add_heading(doc, "一、第一次用？先看这三条", 1)
    add_bullet(doc, "在微信里打开本小程序后，用底部三个按钮就能用到主要功能。")
    add_bullet(
        doc,
        "部分功能需要您用手机号登录（按页面提示授权即可）。例如：完整走完「创建 IP 项目」、使用「我的灵感」等。"
        "若暂时不想登录，部分页面可选择「先去逛逛」，先浏览首页。"
    )
    add_bullet(
        doc,
        "若您已是会员，可在「我的」里查看「算力」并充值；具体权益以「开通会员」页面展示为准。"
    )

    add_heading(doc, "二、「发现」页能做什么", 1)
    add_body(doc, "点底部「发现」，进入首页。")
    add_table_two_col(
        doc,
        [
            ("顶部图片轮播", "展示活动图、品牌介绍等，会随运营更新。"),
            ("通知条", "有重要通知时会在上方显示，请留意。"),
            (
                "功能快捷入口",
                "常见入口包括：脚本洗稿、图文封面、爆款标题、全能对话等（具体以您手机上显示为准）。"
                "点进去一般会打开「AI 对话」，由 AI 助手按该入口帮您处理。",
            ),
            ("运营干货", "可浏览文章列表，点标题看正文。"),
            ("最近落地", "展示近期案例或动态，以页面实际内容为准。"),
        ],
    )
    add_body(doc, "在「发现」页右上角可将小程序分享给微信好友或分享到朋友圈。")

    add_heading(doc, "三、「IP 操作台」怎么用", 1)
    add_body(
        doc,
        "「IP」可以理解为您在短视频/自媒体上的人设与风格。"
        "在这里按「项目」管理：每个项目有一套独立的人设说明，方便您针对不同账号或选题分开操盘。"
    )

    add_heading(doc, "3.1 还没有项目时", 2)
    add_body(doc, "页面会提示您创建项目。按提示填写「IP 信息定位」问卷（分几步完成），用于生成您的定位说明。")

    add_heading(doc, "3.2 已有多个项目时", 2)
    add_body(doc, "会先看项目列表，请点选当前要用的那一个。之后也可在页面顶部等处「切换项目」。")

    add_heading(doc, "3.3 已经选好项目后（主操作界面）", 2)
    add_table_two_col(
        doc,
        [
            ("顶部", "显示当前项目名称、您的信息与积分等；可切换项目。"),
            ("人设卡片", "展示当前项目的人设摘要；点进去可「微调人设」（与新建时步骤相同，保存后更新本项目）。"),
            (
                "今天拍点啥",
                "一块块分类入口，由后台配置。点某一类后，通常会进入「AI 文案生成」帮您写稿，或跳到其他页面（如热点榜）。"
                "若提示「功能开发中」，说明该入口尚在升级，请稍后再试。",
            ),
            ("快捷指令库", "常用指令卡片，点进去进入「AI 文案生成」，并带上适合该指令的提示，方便您直接开写。"),
            ("历史对话", "列出与创作相关的聊天记录；点某一条可接着聊。"),
            ("在列表处向下拉", "可刷新对话列表和快捷入口。"),
            (
                "右下角加号按钮",
                "打开「灵感捕捉」：把突然想到的点子记下来，可加标签，并会保存到「我的灵感」（可与当前项目关联）。"
                "若语音功能提示即将上线，以实际版本为准。",
            ),
        ],
    )

    add_heading(doc, "3.4 新建项目时要填什么", 2)
    add_body(doc, "新建时一般会分几步，例如：")
    add_bullet(doc, "设定身份：行业、角色等，让别人知道「你是谁」。")
    add_bullet(doc, "注入灵魂：价值观、故事感，让人设更立体。")
    add_bullet(doc, "定义风格：语气、说法习惯，让内容口吻统一。")
    add_bullet(doc, "激活大脑：核对信息并确认。")
    add_body(
        doc,
        "填写中途若退出，再次进入时可能询问是否接着上次填写或清空重来。"
        "全部填完后，新建流程会生成「IP 定位报告」，您可在报告页保存为正式项目。"
        "若是「改人设」，最后一步会保存更新并返回，不一定会再走报告页。"
    )
    add_body(doc, "注意：未登录也可能进入填写页，但要完整保存到云端，一般需要先登录。")

    add_heading(doc, "3.5 还可能打开哪些页面", 2)
    add_bullet(doc, "IP 定位报告：查看系统根据问卷生成的定位说明，并保存项目。")
    add_bullet(doc, "AI 文案生成：像聊天一样让助手帮您写文案，可带上分类、快捷指令或历史会话里的上下文。")
    add_bullet(doc, "AI 对话：在「发现」里「全能对话」等入口常用，偏通用对话；与「文案生成」分工以您手机上的实际效果为准。")
    add_bullet(doc, "历史对话（独立列表）：集中查看会话，入口以页面为准。")
    add_bullet(doc, "抖音热点榜单：看热门话题，可用「蹭热点」等与创作结合。")

    add_heading(doc, "四、「我的」里有什么", 1)

    add_heading(doc, "4.1 账号与头像", 2)
    add_bullet(doc, "未登录时点「登录」，按提示完成手机号快捷登录，并需勾选同意用户协议与隐私政策。")
    add_bullet(doc, "登录后可更换头像（按微信提示选择头像并上传）。")
    add_bullet(doc, "手机号在中间四位会打星号显示，保护隐私。")
    add_bullet(doc, "会员用户会显示等级、到期时间等信息（以页面为准）。")

    add_heading(doc, "4.2 会员与算力", 2)
    add_bullet(doc, "普通用户会看到「开通会员」入口，点进去了解权益与购买方式。")
    add_bullet(doc, "已是会员时，可看到「我的算力」余额，并可查看明细、去充值。")

    add_heading(doc, "4.3 菜单说明", 2)
    add_table_two_col(
        doc,
        [
            ("我的灵感", "需登录。查看、搜索您保存的灵感，可归档、置顶；也可从灵感继续生成或对话（以按钮为准）。"),
            ("我要推荐", "邀请好友等活动，按页面说明参与。"),
            ("联系客服", "咨询会员、使用问题等。"),
        ],
    )
    add_body(doc, "在「我的」页面向下拉，可刷新您的账号信息（算力、会员状态等）。")

    add_heading(doc, "五、其他您可能看到的页面", 1)
    add_bullet(doc, "火源文案智能体：智能体介绍或入口。")
    add_bullet(doc, "扫码登录：与电脑网页等端配合登录时使用，按屏幕提示操作。")
    add_bullet(doc, "完善资料：登录后若需要补充资料会出现。")
    add_bullet(doc, "用户协议、隐私政策：请仔细阅读。")
    add_bullet(doc, "文章列表与详情：阅读运营发布的文章。")

    add_heading(doc, "六、常见问题", 1)
    add_body(doc, "1. 提示「功能开发中，请耐心等待」：该菜单或能力尚未开放，或正在升级，请过段时间再试。")
    add_body(
        doc,
        "2.「发现」和「IP 操作台」里都有对话类功能，有什么区别？"
        "——「发现」里多为通用「AI 对话」；「IP 操作台」里多与当前项目、指令、历史会话绑定，更适合围绕您设定的人设写文案。"
    )
    add_body(doc, "3. 加载失败：操作台若加载失败，请点「重试」；仍不行请检查手机网络或稍后再试。")
    add_body(doc, "4. 算力怎么扣：是否扣算力、扣多少，以系统规则和「我的算力」「充值」页说明为准。")

    add_heading(doc, "七、说明", 1)
    add_body(doc, "小程序版本以微信里实际显示为准。首页图片、通知、入口名称等会随运营更新，请以您打开时的界面为准。")
    add_body(doc, "若产品与本文描述不一致，以线上功能为准；欢迎通过「联系客服」反馈。")

    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    p = build()
    print(f"已生成: {p}")
